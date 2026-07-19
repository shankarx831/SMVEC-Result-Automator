import os
import requests
import openpyxl
from docx import Document
from docx.shared import Inches
import card_generator

def test_student():
    reg_no = "25UIT042"
    name = "HARIKANTH R"
    dob = "3/11/2004"
    
    rows = [
        ['1', 'U20HSO101', 'Engineering Mathematics I', '9', 'A', 'PASS'],
        ['1', 'U20ITC102', 'Problem Solving and Python Programming', '9', 'A', 'PASS'],
        ['1', 'U20ITE103', 'Engineering Physics', '10', 'S', 'PASS'],
        ['1', 'U20ITM104', 'Basic Electrical and Electronics Engineering', '9', 'A', 'PASS'],
        ['1', 'U20ITP105', 'Python Programming Laboratory', '10', 'S', 'PASS'],
        ['1', 'U20ITP106', 'Engineering Physics Laboratory', '10', 'S', 'PASS']
    ]
    sgpa = "9.50"
    meta_info = {
        "exam_title": "End Semester Examinations",
        "branch": "B.TECH-Information Technology",
        "announcement_date": "Result Announcement Date: 18/07/2026"
    }

    display_name = name
    
    # Create Excel
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["Name", "Register Number", "DOB"])
    ws.append([display_name, reg_no, dob])
    
    result_data = {}
    for r_col in rows:
        if len(r_col) >= 6:
            sub_name = r_col[2]
            grade_val = r_col[3] # Point
            result_data[sub_name] = grade_val
        elif len(r_col) >= 2:
            result_data[r_col[0]] = r_col[1]

    col_mark = 4
    for sub, grade in result_data.items():
        ws.cell(row=1, column=col_mark).value = sub
        ws.cell(row=2, column=col_mark).value = grade
        col_mark += 1

    ws.cell(row=1, column=col_mark).value = "SGPA"
    ws.cell(row=2, column=col_mark).value = sgpa
    
    excel_path = f"test_{reg_no}_Results.xlsx"
    wb.save(excel_path)
    print(f"Saved Excel: {excel_path}")
    
    # Create Word Doc
    doc = Document()
    img_path = f"test_{reg_no}_card.png"
    card_generator.draw_result_card(reg_no, display_name, meta_info, rows, sgpa, img_path)
    
    doc.add_heading(f'{reg_no} - {display_name}', level=4)
    doc.add_picture(img_path, width=Inches(6.0))
    doc_path = f"test_{reg_no}_Results.docx"
    doc.save(doc_path)
    print(f"Saved Word Doc: {doc_path}")
    
    # Keep the image so we can show it in artifacts
    print(f"Saved Image: {img_path}")

if __name__ == "__main__":
    test_student()
