import os
import sys
import subprocess

# Auto-create and run inside a virtual environment if executed globally (PEP 668 bypass)
def ensure_venv():
    if sys.prefix == sys.base_prefix:
        venv_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "venv")
        if not os.path.exists(venv_dir):
            print("Creating virtual environment...")
            subprocess.check_call([sys.executable, "-m", "venv", venv_dir])
        
        if os.name == 'nt':
            venv_python = os.path.join(venv_dir, "Scripts", "python.exe")
        else:
            venv_python = os.path.join(venv_dir, "bin", "python")
        
        print("Re-starting app within virtual environment...")
        os.execv(venv_python, [venv_python] + sys.argv)

ensure_venv()

# Auto-install dependencies if any required libraries are missing
try:
    import flask
    import openpyxl
    import docx
    import PIL
    import dotenv
    import requests
    import bs4
except ImportError:
    print("Dependencies not met. Automatically installing required packages...")
    try:
        req_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'requirements.txt')
        if os.path.exists(req_path):
            subprocess.check_call([sys.executable, "-m", "pip", "install", "-r", req_path])
        else:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "Flask", "openpyxl", "python-docx", "pillow", "python-dotenv", "requests", "beautifulsoup4", "gunicorn"])
        print("Dependencies successfully installed!")
    except Exception as e:
        print(f"Failed to automatically install dependencies: {e}")
        print("Please run: pip install -r requirements.txt manually.")
        sys.exit(1)

import time
import openpyxl
from flask import Flask, render_template, request, jsonify, send_from_directory, redirect, url_for, session
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches
from PIL import Image
from functools import wraps
from werkzeug.security import check_password_hash
import db
from dotenv import load_dotenv
import functools
import re
import urllib3
import requests
from datetime import datetime, timedelta
import card_generator

# Load configurations from .env file
load_dotenv()

# Initialize SQLite database on app load
db.init_db()

app = Flask(__name__)
# The secret key is needed to keep the client-side sessions secure.
app.secret_key = os.environ.get('FLASK_SECRET_KEY', 'default-dev-key')
app.config['UPLOAD_FOLDER'] = os.path.join(app.root_path, 'uploads')
app.config['RESULTS_FOLDER'] = os.path.join(app.root_path, 'static', 'results')
app.config['MAX_CONTENT_LENGTH'] = 5 * 1024 * 1024  # Max 5MB uploads

# Custom Jinja Filter for converting UTC from SQLite to IST (Indian Standard Time)
@app.template_filter('to_ist')
def to_ist_filter(utc_time_str):
    if not utc_time_str:
        return ""
    try:
        utc_time = datetime.strptime(utc_time_str, '%Y-%m-%d %H:%M:%S')
        ist_time = utc_time + timedelta(hours=5, minutes=30)
        return ist_time.strftime('%d %b %Y, %I:%M %p')
    except Exception:
        return utc_time_str

# Suppress insecure request warnings if fetching without SSL verification
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
ALLOWED_EXTENSIONS = {'xlsx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Ensure directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['RESULTS_FOLDER'], exist_ok=True)

# Global status tracking variable
processing_status = "Idle"

# Auth protection decorators
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session or session.get('role') != 'admin':
            return redirect(url_for('index'))
        return f(*args, **kwargs)
    return decorated_function

def convert_date(date_str):
    if not date_str:
        return ""
    # Support formats like DD.MM.YYYY or DD/MM/YYYY or datetime objects
    date_str = str(date_str).strip()
    if "." in date_str:
        parts = date_str.split(".")
        if len(parts) == 3:
            return f"{parts[0]}/{parts[1]}/{parts[2]}"
    return date_str

@app.route('/')
@login_required
def index():
    raw_history = db.get_user_history(session.get('user_id'))
    history = []
    for h in raw_history:
        # Convert bytes to human readable format
        size_bytes = h['filesize']
        if size_bytes >= 1024 * 1024:
            size_str = f"{size_bytes / (1024 * 1024):.2f} MB"
        else:
            size_str = f"{size_bytes / 1024:.2f} KB"
            
        history.append({
            'filename': h['filename'],
            'size': size_str,
            'file_type': h['file_type'],
            'timestamp': h['timestamp']
        })
    return render_template('index.html', username=session.get('username'), role=session.get('role'), history=history)

@app.route('/login', methods=['GET', 'POST'])
def login():
    if 'user_id' in session:
        return redirect(url_for('index'))
    error = None
    if request.method == 'POST':
        username = request.form.get('username', '').strip().lower()
        password = request.form.get('password', '')
        user = db.get_user(username)
        if user and check_password_hash(user['password_hash'], password):
            session['user_id'] = user['id']
            session['username'] = user['username']
            session['role'] = user['role']
            return redirect(url_for('index'))
        else:
            error = "Invalid username or password."
    return render_template('login.html', error=error)

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

@app.route('/admin')
@admin_required
def admin():
    users = db.get_all_users()
    return render_template('admin.html', users=users, current_username=session.get('username'))

@app.route('/admin/create', methods=['POST'])
@admin_required
def admin_create_user():
    username = request.form.get('username', '').strip().lower()
    password = request.form.get('password', '')
    role = request.form.get('role', 'user')

    if len(password) < 6:
        users = db.get_all_users()
        return render_template('admin.html', users=users, current_username=session.get('username'), error="Password must be at least 6 characters.")

    if db.create_user(username, password, role):
        users = db.get_all_users()
        return render_template('admin.html', users=users, current_username=session.get('username'), success=f"User '{username}' created successfully.")
    else:
        users = db.get_all_users()
        return render_template('admin.html', users=users, current_username=session.get('username'), error=f"User '{username}' already exists.")

@app.route('/admin/delete/<int:user_id>', methods=['POST'])
@admin_required
def admin_delete_user(user_id):
    users = db.get_all_users()
    # Find user details
    target_user = None
    for u in users:
        if u['id'] == user_id:
            target_user = u
            break
            
    if not target_user:
        return render_template('admin.html', users=users, current_username=session.get('username'), error="User not found.")

    if target_user['username'] == 'admin' or target_user['username'] == session.get('username'):
        return render_template('admin.html', users=users, current_username=session.get('username'), error="Cannot delete system admin or self.")

    db.delete_user(user_id)
    # Re-fetch users
    users = db.get_all_users()
    return render_template('admin.html', users=users, current_username=session.get('username'), success=f"User '{target_user['username']}' deleted successfully.")

@app.route('/admin/reset_password/<int:user_id>', methods=['POST'])
@admin_required
def admin_reset_password(user_id):
    new_password = request.form.get('new_password', '')
    users = db.get_all_users()
    
    if len(new_password) < 6:
        return render_template('admin.html', users=users, current_username=session.get('username'), error="Password must be at least 6 characters.")
        
    target_user = None
    for u in users:
        if u['id'] == user_id:
            target_user = u
            break
            
    if not target_user:
        return render_template('admin.html', users=users, current_username=session.get('username'), error="User not found.")
        
    db.update_password(user_id, new_password)
    users = db.get_all_users()
    return render_template('admin.html', users=users, current_username=session.get('username'), success=f"Password for '{target_user['username']}' updated successfully.")

@app.route('/status')
@login_required
def status():
    return jsonify({"status": processing_status})

@app.route('/generate', methods=['POST'])
@login_required
def generate():
    global processing_status
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
        
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
        
    if not allowed_file(file.filename):
        return jsonify({"error": "Invalid file type. Only .xlsx files are allowed."}), 400

    portal_url = request.form.get('url', 'http://exam.smvec.ac.in/exam_result_ug_pg_apr2026_regular/').strip()
    gen_excel = request.form.get('gen_excel') == 'true'
    gen_word = request.form.get('gen_word') == 'true'

    # Save uploaded template
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)

    processing_status = "Loading Excel workbook..."
    try:
        workbook = openpyxl.load_workbook(filepath)
        worksheet = workbook.active
    except Exception as e:
        return jsonify({"error": f"Failed to read Excel file: {str(e)}"}), 400

    # Start pure HTTP session setup (no Selenium or Chrome needed!)
    processing_status = "Connecting to SMVEC Result Portal..."
    http_session = requests.Session()
    submit_btn_info = None

    # Output filenames
    base_name = os.path.splitext(filename)[0]
    out_excel_name = f"{base_name}_Results.xlsx"
    out_word_name = f"{base_name}_Results.docx"
    out_excel_path = os.path.join(app.config['RESULTS_FOLDER'], out_excel_name)
    out_word_path = os.path.join(app.config['RESULTS_FOLDER'], out_word_name)

    doc = Document() if gen_word else None

    idx = 2
    num_rows = worksheet.max_row
    
    # Auto-detect columns
    col_name = 1
    col_reg = 2
    col_dob = 3
    for c in range(1, worksheet.max_column + 1):
        header = str(worksheet.cell(row=1, column=c).value).lower()
        if 'name' in header:
            col_name = c
        elif 'register' in header or 'reg' in header:
            col_reg = c
        elif 'dob' in header or 'date of birth' in header or 'birth' in header:
            col_dob = c
            
    # Calculate actual number of valid student rows (avoids openpyxl ghost rows issue)
    actual_student_count = 0
    for r in range(2, num_rows + 1):
        if worksheet.cell(row=r, column=col_name).value or worksheet.cell(row=r, column=col_reg).value:
            actual_student_count += 1
            
    # Track dynamic subjects to align columns perfectly
    subject_cols = {}
    current_max_col = worksheet.max_column
    sgpa_tracker = {}
    
    processed_count = 0
    
    try:
        while idx <= num_rows:
            name = worksheet.cell(row=idx, column=col_name).value
            reg_no = worksheet.cell(row=idx, column=col_reg).value
            dob = worksheet.cell(row=idx, column=col_dob).value

            if not name and not reg_no:
                idx += 1
                continue
                
            processed_count += 1
            reg_no = str(reg_no).strip()
            dob_formatted = convert_date(dob)
            processing_status = f"Processing {processed_count}/{actual_student_count}: {name or ''} ({reg_no})"

            try:
                rows, sgpa, student_name, meta_info, err = card_generator.fetch_and_parse_result(
                    http_session, portal_url, reg_no, dob_formatted, submit_btn_info
                )
                if meta_info and meta_info.get("submit_btn_info"):
                    submit_btn_info = meta_info.get("submit_btn_info")

                if err:
                    print(f"Skipping record row {idx} ({reg_no}): {err}")
                else:
                    display_name = student_name or name or "Student"

                    # 1. Extract grades if Excel option is selected
                    if gen_excel and rows:
                        for r_col in rows:
                            if len(r_col) >= 6:
                                sub_name = r_col[2]
                                grade_val = r_col[4] # Index 4 is the Letter Grade (e.g. A, S)
                                
                                if sub_name not in subject_cols:
                                    current_max_col += 1
                                    subject_cols[sub_name] = current_max_col
                                    worksheet.cell(row=1, column=current_max_col).value = sub_name
                                    
                                worksheet.cell(row=idx, column=subject_cols[sub_name]).value = grade_val
                                
                        if sgpa:
                            sgpa_tracker[idx] = sgpa

                    # 2. Take exact result card picture and add to Word doc if selected
                    if gen_word and rows:
                        temp_img_path = os.path.join(app.config['UPLOAD_FOLDER'], f"card_{reg_no}.png")
                        card_generator.draw_result_card(reg_no, display_name, meta_info, rows, sgpa, temp_img_path)

                        doc.add_heading(f'{reg_no} {display_name}', level=4)
                        p_pic = doc.add_picture(temp_img_path, width=Inches(6.0))
                        
                        # Add page break to the end of the picture paragraph instead of a new paragraph
                        # to avoid empty lines at the top of the next page
                        from docx.enum.text import WD_BREAK
                        doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)

                        if os.path.exists(temp_img_path):
                            os.remove(temp_img_path)

            except Exception as inner_err:
                print(f"Error processing record row {idx} ({reg_no}): {inner_err}")

            idx += 1

        # Remove the very first empty paragraph created by default if it's empty
        if gen_word and len(doc.paragraphs) > 0 and doc.paragraphs[0].text.strip() == "":
            p = doc.paragraphs[0]
            p._element.getparent().remove(p._element)

        # Write SGPA column dynamically at the very end
        if gen_excel and sgpa_tracker:
            sgpa_col_idx = current_max_col + 1
            worksheet.cell(row=1, column=sgpa_col_idx).value = "SGPA"
            for row_idx, sgpa_val in sgpa_tracker.items():
                worksheet.cell(row=row_idx, column=sgpa_col_idx).value = sgpa_val

        # Save output assets
        if gen_excel:
            workbook.save(out_excel_path)
            try:
                size = os.path.getsize(out_excel_path)
                db.add_history_record(session.get('user_id'), out_excel_name, out_excel_path, size, 'excel')
            except Exception as e:
                print(f"Failed to save Excel history: {e}")
                
        if gen_word:
            doc.save(out_word_path)
            try:
                size = os.path.getsize(out_word_path)
                db.add_history_record(session.get('user_id'), out_word_name, out_word_path, size, 'word')
            except Exception as e:
                print(f"Failed to save Word history: {e}")

        # Enforce 50MB user quota restriction
        try:
            db.enforce_user_quota(session.get('user_id'))
        except Exception as e:
            print(f"Failed to enforce user quota: {e}")

    finally:
        processing_status = "Idle"

    # Increment metrics in DB
    try:
        db.increment_generation_count(
            session.get('username'),
            excel_inc=1 if gen_excel else 0,
            word_inc=1 if gen_word else 0
        )
    except Exception as e:
        print(f"Failed to increment user metrics: {e}")

    return jsonify({
        "success": True,
        "excel_url": f"/download/{out_excel_name}" if gen_excel else None,
        "word_url": f"/download/{out_word_name}" if gen_word else None
    })

@app.route('/download_sample')
@login_required
def download_sample():
    return send_from_directory(os.path.join(app.root_path, 'data'), '21-25 IT A.xlsx', as_attachment=True, download_name='sample_format.xlsx')

@app.route('/download/<filename>')
@login_required
def download_file(filename):

    return send_from_directory(app.config['RESULTS_FOLDER'], filename, as_attachment=True)

if __name__ == '__main__':
    import webbrowser
    from threading import Timer

    port = int(os.environ.get('FLASK_PORT', 5001))

    def open_browser():
        webbrowser.open_new(f"http://127.0.0.1:{port}")

    # Only open browser on the main thread, not on reloader threads
    if not os.environ.get("WERKZEUG_RUN_MAIN"):
        Timer(1.5, open_browser).start()

    app.run(host='0.0.0.0', port=port, debug=False)
